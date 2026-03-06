#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
从 DOCX 文档中提取标题（title）和脚注（footnote）
依据：文档开头的目录页（TOC）

作者：Lingma
"""

import os
import pandas as pd
import argparse
import sys
from docx import Document
from typing import List, Dict, Optional

# 导入各个功能模块（使用绝对导入）
from config import DEFAULT_INPUT_FILE, TABLE_NUMBER_PATTERN
from toc_extractor import extract_toc_titles, extract_full_toc
from title_matcher import find_title_paragraph_index
from table_locator import find_next_table_after_index
from output_name_generator import generate_output_name
from data_processor import process_title_split, process_footnote_split, process_footnote_encoding_version, check_xxx_patterns, reorder_columns
from utils import process_superscript_subscript_text, is_fully_italic, contains_programming_keyword, extract_content_before_keyword

# 导入增强版脚注提取器
from enhanced_footnote_extractor import extract_footnotes_in_range

# GUI相关导入
try:
    import tkinter as tk
    from gui import ExtractTitleFootnoteGUI
    GUI_AVAILABLE = True
except ImportError:
    GUI_AVAILABLE = False

def process_document(input_file: str, output_file: str, max_footnote_cols: int = 7, workspace: str = None, study_number: str = None, custom_keywords: List[str] = None, cancellation_token = None, generate_shell_doc: bool = True):
    """处理DOCX文档的主要函数
    
    Args:
        input_file: 输入的DOCX文件路径
        output_file: 输出的Excel文件路径
        max_footnote_cols: 最大脚注列数，默认为7
        workspace: 工作目录路径，默认为None时从input_file推导
        study_number: 项目编号，用于文件名前缀
        custom_keywords: 自定义关键词列表，用于脚注终止条件
        cancellation_token: 取消令牌，用于检查是否需要取消处理
    """
    
    # 如果没有提供workspace，则从input_file推导
    if workspace is None:
        workspace = os.path.dirname(input_file)
    
    # 检查输入文件
    if not os.path.isfile(input_file):
        print(f"错误：未找到输入文件 {input_file}。")
        print("请提供正确的DOCX文件路径，例如：")
        print("  python extract_titles_footnotes.py \"D:\\path\\to\\your\\file.docx\"")
        return False
    
    try:
        doc = Document(input_file)
    except Exception as e:
        print(f"错误：无法打开文档 {input_file} — {e}")
        print("提示：请关闭 Word 中该文件后再试。")
        return False
    
    paragraphs = doc.paragraphs
    tables = doc.tables
    
    # 步骤1: 提取目录标题（返回原始+清洗后）
    # 提取目录中的标题（两个版本：完整目录用于章节名匹配，有效标题用于主要处理）
    toc_entries = extract_toc_titles(paragraphs)  # 原有的有效标题（表/图/列表开头）
    full_toc_entries = extract_full_toc(paragraphs)  # 新增：完整目录（包含所有类型标题）
  
    
    # 分别获取清洗后的标题列表
    toc_cleaned_list = [entry["cleaned"] for entry in toc_entries]
    full_toc_cleaned_list = [entry["cleaned"] for entry in full_toc_entries]
    toc_index_list = [entry["idx_num"] for entry in toc_entries]  # 目录标题在段落中的索引位置
    
    #print(f"提取到 {len(toc_entries)} 个有效目录标题（表/图/列表开头）")
    #print(f"提取到 {len(full_toc_entries)} 个完整目录标题（包含所有类型）")
    
    # 显示一些示例
    #if full_toc_entries:
    #    print("\n完整目录标题示例（前5个）:")
    #    for i, entry in enumerate(full_toc_entries[:5]):
    #        print(f"  {i+1}. {entry['cleaned']}")
    
    #if toc_entries:
    #    print("\n有效目录标题示例（前5个）:")
    #    for i, entry in enumerate(toc_entries[:5]):
    #        print(f"  {i+1}. {entry['cleaned']}")
    
    if not toc_entries:
        print("警告：未在文档中找到任何有效目录标题。")
        return True
    
    toc_raw_list = [entry["raw"] for entry in toc_entries]
    #toc_titles_set = set(toc_cleaned_list)
    
    results = []
    
    # 步骤2: 对每个目录标题处理
    for idx, (raw_title, cleaned_title, toc_index) in enumerate(zip(toc_raw_list, toc_cleaned_list, toc_index_list), 1):
        # 检查是否需要取消
        if cancellation_token and cancellation_token.is_cancelled():
            print("处理被用户取消")
            return False
        
        print(f"\n{'='*50}")
        print(f"处理第 {idx} 个表格: {cleaned_title}")
        print(f"{'='*50}")
        
        # 2.1 查找正文标题段落
        para_idx = toc_index
        matched_para_text = raw_title if para_idx is not None else ""
        
        # 处理正文标题的上标下标
        processed_title_text = ""
        if para_idx is not None:
            title_paragraph = paragraphs[para_idx]
            processed_title_text = process_superscript_subscript_text(title_paragraph)
        
        if para_idx is None:
            output_name = generate_output_name(cleaned_title)
            results.append({
                "序号": idx,
                "目录标题": cleaned_title,
                "Output Name": output_name,
                "正文标题": "",
                "processed_title": "",  # 新增列：处理后的标题
                "匹配状态": "未匹配",
                "表格索引": None,
                "脚注": "",
                "processed_footnote": "",  # 新增列
                "状态": "标题未匹配",
                "备注": f"{cleaned_title} 正文没有匹配的表格，请更新目录"
            })
            continue
        
        # 2.2 关联表格
        table_idx = find_next_table_after_index(tables, para_idx, paragraphs)
        if table_idx is None:
            output_name = generate_output_name(cleaned_title)
            results.append({
                "序号": idx,
                "目录标题": cleaned_title,
                "Output Name": output_name,
                "正文标题": matched_para_text,
                "processed_title": processed_title_text,  # 新增列
                "匹配状态": "匹配",
                "表格索引": None,
                "脚注": "",
                "processed_footnote": "",  # 新增列
                "状态": "无关联表格",
                "备注": f"{cleaned_title} 后无表格"
            })
            continue
        
        # 2.3 使用增强版脚注提取器提取脚注（三级策略）
        # 确定下一个标题的位置作为范围边界
        next_title_idx = None
        if idx < len(toc_raw_list):
            next_raw_title = toc_raw_list[idx]
            next_para_idx = find_title_paragraph_index(paragraphs, next_raw_title)
            if next_para_idx is not None:
                next_title_idx = next_para_idx
        
        #print(f"当前标题段落索引: {para_idx}")
        #print(f"关联表格索引: {table_idx}")
        #print(f"下一个标题索引: {next_title_idx}")
        
        # 调用增强版脚注提取器（传入完整目录用于章节名匹配）
        footnote_text, processed_footnote_text, footnote_status = extract_footnotes_in_range(
            doc, para_idx, next_title_idx, full_toc_cleaned_list, custom_keywords, cancellation_token
        )
        
        # 检查是否在脚注提取过程中被取消
        if cancellation_token and cancellation_token.is_cancelled():
            print("处理被用户取消")
            return False
        
        # 记录结果
        output_name = generate_output_name(cleaned_title)
        results.append({
            "序号": idx,
            "目录标题": cleaned_title,
            "Output Name": output_name,
            "正文标题": matched_para_text,
            "processed_title": processed_title_text,  # 新增列
            "匹配状态": "匹配",
            "表格索引": table_idx,
            "脚注": footnote_text,
            "processed_footnote": processed_footnote_text,  # 新增列
            "状态": footnote_status,
            "备注": ""
        })
    
    # 检查是否需要取消（在后续处理之前）
    if cancellation_token and cancellation_token.is_cancelled():
        print("处理被用户取消")
        return False
    
    # 步骤3: 处理正文标题拆分
    results, max_title_lines = process_title_split(results)
    
    # 步骤4: 处理编码版本替换（在脚注拆分前）
    results = process_footnote_encoding_version(results)
    
    # 步骤5: 检查XXX模式（在脚注拆分前）
    results = check_xxx_patterns(results)
    
    # 步骤6: 处理脚注拆分（基于更新后的processed_footnote，支持最大列数限制）
    results, max_footnote_lines = process_footnote_split(results, max_footnote_cols)
    
    # 创建DataFrame
    df = pd.DataFrame(results)
    
    # 重命名列
    column_mapping = {
        '目录标题': '标题',
        '状态': '脚注匹配状态'
    }
    df = df.rename(columns=column_mapping)
    
    # 在Output Name前新增四列：Batch、Type、TLF、Prgm Name
    import re
    
    # 为每行添加新列
    new_columns_data = []
    for _, row in df.iterrows():
        title = row.get('标题', '')
        output_name = row.get('Output Name', '')
        
        # Batch列：所有值都是1
        batch_value = 1
        
        # Type列：根据标题前缀判断
        if title.startswith('表'):
            type_value = 'TLF-Table'
        elif title.startswith('列表'):
            type_value = 'TLF-Listing'
        elif title.startswith('图'):
            type_value = 'TLF-Figure'
        else:
            type_value = ''  # 默认空值
            
        # TLF列：提取编号并替换前缀
        tlf_value = ''
        if title:
            # 使用正则表达式提取编号部分（表/列表/图后面的编号）
            pattern = r'^(表|列表|图)\s*([^\s]+)'
            match = re.match(pattern, title)
            if match:
                prefix = match.group(1)
                number = match.group(2)
                # 根据前缀转换为T/L/F
                prefix_map = {'表': 'T', '列表': 'L', '图': 'F'}
                converted_prefix = prefix_map.get(prefix, '')
                tlf_value = f"{converted_prefix}{number}"
        
        # Prgm Name列：等于Output Name
        prgm_name_value = output_name
        
        new_columns_data.append({
            'Batch': batch_value,
            'Type': type_value,
            'TLF': tlf_value,
            'Prgm Name': prgm_name_value
        })
    
    # 将新列数据转换为DataFrame并合并
    new_columns_df = pd.DataFrame(new_columns_data)
    df = pd.concat([new_columns_df, df], axis=1)
    
    # 生成包含超链接的辅助 Excel 文件
    # 创建辅助 DataFrame，包含文件名称、超链接、表格名称、备注四列
    auxiliary_data = []
    for _, row in df.iterrows():
        file_name = row.get('Output Name', '')
        table_name = row.get('标题', '')
        processed_title_val = row.get('processed_title', '')
        processed_footnote_val = row.get('processed_footnote', '')
        # 创建超链接公式，指向同目录下的 .rtf 文件
        #hyperlink_formula = f'=HYPERLINK("{file_name}.rtf", "{file_name}")'
        auxiliary_data.append({
            '文件名称': file_name,
            '表格名称': table_name,
            'processed_title': processed_title_val,
            'processed_footnote': processed_footnote_val,
            '备注': ''
        })
    
    # 创建辅助DataFrame
    #auxiliary_df = pd.DataFrame(auxiliary_data)
    
    # 删除不需要输出到Excel的列
    columns_to_drop = ['正文标题', '匹配状态', 'processed_title', '表格索引', 'processed_footnote']
    existing_columns_to_drop = [col for col in columns_to_drop if col in df.columns]
    if existing_columns_to_drop:
        df = df.drop(columns=existing_columns_to_drop)
    
    # 调整列顺序
    df = reorder_columns(df, max_title_lines, max_footnote_lines)
    
    try:
        # 保存主结果文件
        df.to_excel(output_file, index=False)
        
        # 生成辅助文件路径 - 根据项目编号确定文件名
        if study_number:
            auxiliary_output_file = os.path.join(workspace, f"{study_number}_LOT.xlsx")
        else:
            auxiliary_output_file = os.path.join(workspace, "LOT.xlsx")
        
        # 使用openpyxl直接创建辅助文件以支持正确的超链接样式
        # 导入openpyxl（如果尚未导入）
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
        
        # 创建工作簿和工作表
        wb = Workbook()
        ws = wb.active
        ws.title = "LOT"
        
        # 定义边框样式（所有边框）
        thin_border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )
        
        # 写入表头
        headers = ['文件名称', '表格名称', '备注']
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num, value=header)
            cell.font = Font(bold=True)
            cell.border = thin_border
            # 设置居中对齐
            cell.alignment = Alignment(horizontal='center')
            
        # 写入数据行
        for row_idx, row_data in enumerate(auxiliary_data, 2):
            file_name = row_data['文件名称']
            table_name = row_data['表格名称']
            remark = row_data['备注']
            
            # 文件名称列：设置为超链接格式
            cell = ws.cell(row=row_idx, column=1, value=file_name)
            # 设置超链接（相对路径）
            cell.hyperlink = f"./{file_name}.rtf"
            # 应用Hyperlink样式
            cell.style = "Hyperlink"
            # 添加边框
            cell.border = thin_border
            
            # 表格名称列
            cell = ws.cell(row=row_idx, column=2, value=table_name)
            cell.border = thin_border
            
            # 备注列
            cell = ws.cell(row=row_idx, column=3, value=remark)
            cell.border = thin_border
            
            # 自动调整列宽
            for col in ws.columns:
                max_length = 0
                column = col[0].column_letter  # Get the column name
                for cell in col:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = max(min(max_length + 2, 50), 10)  # Max width 50, minimum width 10
                
                # Special handling for '表格名称' column (column B)
                if column == 'B':  # '表格名称' is the second column
                    # Set minimum width of 30 for table name column to accommodate longer names
                    adjusted_width = max(adjusted_width, 100)
                
                ws.column_dimensions[column].width = adjusted_width
            
            # 保存文件
            wb.save(auxiliary_output_file)
        
        successful_count = sum(1 for r in results if r['状态'] == '成功')
        print(f"✅ 提取完成！结果已保存至：{output_file}")
        print(f"✅ 辅助文件已保存至：{auxiliary_output_file}")
        print(f"共处理 {len(toc_entries)} 个表格，其中 {successful_count} 项脚注提取成功。")
        
        # 标题行数信息
        if max_title_lines > 1:
            print(f"检测到正文标题最多包含 {max_title_lines} 行，已自动拆分为 {max_title_lines} 个独立的title列。")
        else:
            print("所有正文标题均为单行，已合并为一个title列。")
            
        # 脚注行数信息
        if max_footnote_lines > 1:
            print(f"检测到脚注最多包含 {max_footnote_lines} 行，已自动拆分为 {max_footnote_lines} 个独立的footnote列。")
        elif max_footnote_lines == 1:
            print("所有脚注均为单行，已合并为一个footnote列。")
        else:
            print("未检测到有效脚注内容。")
        
        # 如果需要生成 shell 文档，则调用生成函数
        if generate_shell_doc:
            print(f"\n📄 正在生成 shell 文档（只包含标题和脚注）...")
            
            # 生成 shell 文档（使用 auxiliary_data，避免重复读取）
            from process_shell_content import process_shell_file
            process_shell_file(
                auxiliary_data=auxiliary_data,
                source_path=input_file,
                project_id=study_number
            )
        
        return True
    except Exception as e:
        print(f"错误：无法写入Excel文件 — {e}")
        return False

def main():
    """程序主入口"""
    parser = argparse.ArgumentParser(description="从DOCX提取标题与脚注（基于目录页）")
    parser.add_argument("input_file", nargs="?", 
                        default=None,  # 改为None，这样可以检测是否提供了参数
                        help="输入的DOCX文件路径")
    parser.add_argument("--max-footnote-cols", type=int, default=7,
                        help="最大footnote列数，默认为7")
    parser.add_argument("--study-number", type=str, default=None,
                        help="项目编号，用于文件名前缀")
    parser.add_argument("--custom-keywords", type=str, nargs='*', default=None,
                        help="自定义脚注关键词列表，例如: --custom-keywords keyword1 keyword2")
    args = parser.parse_args()
    
    # 如果没有提供输入文件参数，启动GUI模式
    if args.input_file is None:
        if not GUI_AVAILABLE:
            print("错误：GUI模块不可用，请安装tkinter或使用命令行模式。")
            print("命令行使用方式：")
            print("  python main.py \"path/to/document.docx\"")
            return
        
        try:
            # 创建根窗口
            root = tk.Tk()
            app = ExtractTitleFootnoteGUI(root)
            print("🚀 GUI界面已启动")
            root.mainloop()
        except Exception as e:
            print(f"GUI启动失败：{e}")
            import traceback
            traceback.print_exc()
        return
    
    # 命令行模式
    input_file = args.input_file
    max_footnote_cols = args.max_footnote_cols
    study_number = args.study_number
    custom_keywords = args.custom_keywords
    workspace = os.path.dirname(input_file)
    
    # 根据是否提供项目编号来确定输出文件名
    if study_number:
        output_file = os.path.join(workspace, f"{study_number}_TF_Contents.xlsx")
        print(f"📋 项目编号: {study_number}")
    else:
        output_file = os.path.join(workspace, "TF_Contents.xlsx")
        print("📋 未指定项目编号，使用默认文件名")
    
    print(f"📋 最大footnote列数设置为: {max_footnote_cols}")
    if custom_keywords:
        print(f"📋 自定义关键词: {custom_keywords}")
    else:
        from config import DEFAULT_CUSTOM_KEYWORDS
        print(f"📋 使用默认关键词: {DEFAULT_CUSTOM_KEYWORDS}")
    process_document(input_file, output_file, max_footnote_cols, workspace, study_number, custom_keywords)

if __name__ == "__main__":
    main()