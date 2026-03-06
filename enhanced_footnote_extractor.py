#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
增强版脚注提取器 - 实现基于内容范围关联的三级策略提取逻辑

更新逻辑说明：
1. 首先通过目录定位到标题后，遍历文档中的所有内容
2. 将每部分内容与其上最近的标题关联起来
3. 在关联的内容范围内查找表格、图片或伪占位段落
4. 根据查找到的对象确定脚注起始位置
5. 确保脚注提取不超出标题关联的范围
"""

import re
from typing import List, Tuple, Optional, Dict, Any
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

# 导入必要的工具函数和常量
try:
    from utils import (
        process_superscript_subscript_text, 
        is_fully_italic, 
        contains_custom_keyword, 
        extract_content_before_custom_keyword,
        extract_content_before_custom_keyword_from_text
    )
    from config import TABLE_NUMBER_PATTERN
except ImportError:
    # 如果相对导入失败，尝试绝对导入
    import sys
    import os
    sys.path.append(os.path.dirname(__file__))
    from utils import (
        process_superscript_subscript_text, 
        is_fully_italic, 
        contains_custom_keyword, 
        extract_content_before_custom_keyword,
        extract_content_before_custom_keyword_from_text
    )
    from config import TABLE_NUMBER_PATTERN

# 定义重复表/图/列表的正则模式（支持空格可选）
REPEATED_PLACEHOLDER_PATTERN = re.compile(r'^\s*(重复)?\s*(表|图|列表)\s*(?:\d+(?:\.\d+)*(?:\.?[a-zA-Z]+)?|略)[。\.]?\s*')
FIGURE_OMITTED_PATTERN = re.compile(r'^\s*图略\s*')

def build_content_ranges(document: DocxDocument, title_indices: List[int]) -> List[Dict[str, Any]]:
    """
    构建内容范围映射：将文档内容按标题分割成不同的范围
    
    Args:
        document: DOCX文档对象
        title_indices: 标题段落索引列表（已排序）
        
    Returns:
        list: 包含每个范围信息的字典列表
    """
    paragraphs = document.paragraphs
    ranges = []
    
    for i, title_idx in enumerate(title_indices):
        # 确定范围结束位置
        if i < len(title_indices) - 1:
            end_idx = title_indices[i + 1]
        else:
            end_idx = len(paragraphs)
        
        range_info = {
            'title_index': title_idx,
            'title_text': paragraphs[title_idx].text.strip(),
            'start_index': title_idx,
            'end_index': end_idx,
            'elements': []  # 存储范围内的所有元素
        }
        
        # 收集此范围内的所有元素（段落、表格和图片）
        collect_elements_in_range(document, range_info)
        ranges.append(range_info)
    
    return ranges

def collect_elements_in_range(document: DocxDocument, range_info: Dict[str, Any]):
    """
    收集指定范围内的所有元素（段落、表格和图片）
    
    Args:
        document: DOCX文档对象
        range_info: 范围信息字典
    """
    paragraphs = document.paragraphs
    tables = document.tables
    
    # 获取文档所有元素的全局顺序
    body = paragraphs[0]._element.getparent()
    all_elements = []
    
    # 定义图片相关的命名空间前缀
    PIC_NAMESPACE = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'
    
    for elem in body:
        if elem.tag.endswith('p'):
            # 检查段落中的图片
            para_obj = None
            for p in paragraphs:
                if p._element is elem:
                    para_obj = p
                    break
            
            if para_obj:
                # 检查段落中的每个run是否包含图片
                has_picture = False
                for run in para_obj.runs:
                    # 使用xpath查找pic:pic元素来检测图片
                    pic_elements = run.element.xpath('.//pic:pic')
                    if pic_elements:
                        has_picture = True
                        # 为每个图片创建单独的元素记录
                        for i, pic_elem in enumerate(pic_elements):
                            all_elements.append(('picture', pic_elem, {
                                'paragraph': elem,
                                'run': run,
                                'picture_index_in_run': i
                            }))
                        break
                
                # 如果段落包含图片，也要添加段落本身
                if has_picture:
                    all_elements.append(('paragraph_with_pictures', elem))
                else:
                    all_elements.append(('paragraph', elem))
                    
        elif elem.tag.endswith('tbl'):
            # 直接添加表格元素，不检查表格中的图片
            table_obj = None
            for t in tables:
                if t._element is elem:
                    table_obj = t
                    break
            
            if table_obj:
                all_elements.append(('table', elem))
    
    start_idx = range_info['start_index']
    end_idx = range_info['end_index']
    
    # 找到范围起始和结束在全局元素中的位置
    start_pos = None
    end_pos = None
    
    # 找起始位置
    start_para_elem = paragraphs[start_idx]._element
    for i, item in enumerate(all_elements):
        elem_type = item[0]
        elem = item[1]
        if (elem_type == 'paragraph' or elem_type == 'paragraph_with_pictures') and elem is start_para_elem:
            start_pos = i
            break
    
    # 找结束位置
    if end_idx < len(paragraphs):
        end_para_elem = paragraphs[end_idx]._element
        for i, item in enumerate(all_elements):
            elem_type = item[0]
            elem = item[1]
            if (elem_type == 'paragraph' or elem_type == 'paragraph_with_pictures') and elem is end_para_elem:
                end_pos = i
                break
    else:
        end_pos = len(all_elements)
    
    if start_pos is None:
        return
    
    # 收集范围内的元素
    elements = []
    for i in range(start_pos, end_pos if end_pos is not None else len(all_elements)):
        item = all_elements[i]
        elem_type = item[0]
        elem = item[1]
        
        if elem_type == 'paragraph' or elem_type == 'paragraph_with_pictures':
            # 找到对应的段落索引
            para_idx = None
            for j, p in enumerate(paragraphs):
                if p._element is elem:
                    para_idx = j
                    break
            if para_idx is not None and start_idx <= para_idx < end_idx:
                elements.append({
                    'type': 'paragraph',
                    'element': elem,
                    'paragraph_index': para_idx,
                    'text': paragraphs[para_idx].text.strip(),
                    'paragraph_object': paragraphs[para_idx]
                })
        elif elem_type == 'table' or elem_type == 'table_with_pictures':
            # 找到对应的表格索引
            table_idx = None
            for j, t in enumerate(tables):
                if t._element is elem:
                    table_idx = j
                    break
            if table_idx is not None:
                elements.append({
                    'type': 'table',
                    'element': elem,
                    'table_index': table_idx,
                    'table_object': tables[table_idx]
                })
        elif elem_type == 'picture':
            # 图片元素处理
            picture_info = item[2] if len(item) > 2 else {}
            elements.append({
                'type': 'picture',
                'element': elem,
                'picture_details': picture_info,
                'picture_index': i,  # 使用全局索引作为图片标识
            })
    
    range_info['elements'] = elements

def is_placeholder_paragraph(paragraph: Paragraph) -> bool:
    """
    判断段落是否为伪占位段落（重复表/图/列表 或 图略）
    
    Args:
        paragraph: 段落对象
        
    Returns:
        bool: 是否为占位段落
    """
    text = paragraph.text.strip()
    
    # 检查重复表/图/列表模式
    if REPEATED_PLACEHOLDER_PATTERN.match(text):
        #print(f"DEBUG: 检测到占位段落: '{text}' - 匹配REPEATED_PLACEHOLDER_PATTERN")
        return True
    
    # 检查图略模式
    if FIGURE_OMITTED_PATTERN.match(text):
        #print(f"DEBUG: 检测到占位段落: '{text}' - 匹配FIGURE_OMITTED_PATTERN")
        return True
    
    return False

def find_footnote_start_position(range_elements: List[Dict], xml_images: Optional[List[Dict]] = None) -> Optional[int]:
    """
    在范围元素中查找脚注起始位置（四级策略）
    
    Args:
        range_elements: 范围内的元素列表
        xml_images: 从XML解析的图片信息列表（可选）
        
    Returns:
        int: 脚注起始段落索引，如果未找到则返回None
    """
    # 第一级：查找真实表格
    for i, element_info in enumerate(range_elements):
        if element_info['type'] == 'table':
            # 找到表格后的第一个段落
            for j in range(i + 1, len(range_elements)):
                next_element = range_elements[j]
                if next_element['type'] == 'paragraph':
                    #print(f"DEBUG: 找到表格，脚注起始于段落索引: {next_element['paragraph_index']}")
                    return next_element['paragraph_index']
    
    # 第二级：查找图片（先尝试python-docx识别的图片）
    for i, element_info in enumerate(range_elements):
        if element_info['type'] == 'picture':
            # 增加调试信息：输出picture元素的详细位置信息
            picture_index = element_info.get('picture_index', 'unknown')
            picture_details = element_info.get('picture_details', {})
            
            # 尝试获取段落索引信息
            paragraph_index = None
            if 'paragraph' in picture_details:
                # 如果图片详情中有段落信息，查找对应的段落元素
                target_para_elem = picture_details['paragraph']
                # 在range_elements中查找对应的段落元素
                for range_elem in range_elements:
                    if (range_elem['type'] == 'paragraph' and 
                        range_elem.get('element') is target_para_elem):
                        paragraph_index = range_elem.get('paragraph_index')
                        break
            
            if paragraph_index is None:
                # 如果无法从详情中获取，尝试从元素本身获取
                paragraph_index = element_info.get('paragraph_index')
            
            #print(f"DEBUG: 检测到图片元素 - 索引: {picture_index}, 段落索引: {paragraph_index}, 详情: {picture_details}")
            
            # 找到图片后的第一个段落
            for j in range(i + 1, len(range_elements)):
                next_element = range_elements[j]
                if next_element['type'] == 'paragraph':
                    #print(f"DEBUG: 找到图片（python-docx），脚注起始于段落索引: {next_element['paragraph_index']}")
                    return next_element['paragraph_index'] + 1 # 图片后的第一个段落
    
    # 第三级：查找伪占位段落
    for element_info in range_elements:
        if element_info['type'] == 'paragraph':
            paragraph = element_info['paragraph_object']
            if is_placeholder_paragraph(paragraph):
                # 找到占位段落后的第一个段落
                para_idx = element_info['paragraph_index']
                #print(f"DEBUG: 找到占位段落，脚注起始于段落索引: {para_idx + 1}")
                return para_idx + 1  # 占位段落后的第一个段落
    
    # 第四级：从标题后第一个段落开始
    for element_info in range_elements:
        if element_info['type'] == 'paragraph':
            # 返回第一个段落（标题后的第一个段落）
            #print(f"DEBUG: 使用默认策略，脚注起始于段落索引: {element_info['paragraph_index']}")
            return element_info['paragraph_index']
    
    return None

def extract_footnotes_in_range(document: DocxDocument, start_index: int,
                              end_index: Optional[int] = None, 
                              toc_cleaned_titles: List[str] = None,
                              custom_keywords: List[str] = None,
                              cancellation_token = None) -> Tuple[str, str, str]:
    """
    在指定范围内提取脚注内容（三级策略）
    
    Args:
        document: DOCX文档对象
        start_index: 当前标题段落索引
        end_index: 下一标题段落索引（可选，默认为文档末尾）
        toc_cleaned_titles: 目录中清洗后的标题列表，用于章节名匹配终止条件
        custom_keywords: 自定义关键词列表，将与默认的programming/programmer一起用于脚注终止条件
        cancellation_token: 取消令牌，用于检查是否需要取消处理
        
    Returns:
        tuple: (原始脚注文本, 处理后的脚注文本, 脚注状态)
    """
    # 检查是否需要取消
    if cancellation_token and cancellation_token.is_cancelled():
        return "", "", "已取消"
    
    paragraphs = document.paragraphs
    
    # 确定范围结束位置
    if end_index is None:
        end_index = len(paragraphs)
    
    # 构建范围信息
    range_info = {
        'title_index': start_index,
        'title_text': paragraphs[start_index].text.strip(),
        'start_index': start_index,
        'end_index': end_index
    }
    
    # 收集范围内的元素
    collect_elements_in_range(document, range_info)
    range_elements = range_info['elements']
    
    if not range_elements:
        return "", "", "脚注未找到"
    
    # 查找脚注起始位置
    start_para_idx = find_footnote_start_position(range_elements)
    
    if start_para_idx is None:
        return "", "", "脚注未找到"
    
    # 确保起始位置在范围内
    start_para_idx = max(start_para_idx, start_index + 1)
    start_para_idx = min(start_para_idx, end_index - 1)
    
    # 提取脚注内容
    footnote_lines = []
    processed_footnote_lines = []
    current_idx = start_para_idx
    max_lines = 50  # 限制最大行数
    found_placeholder = False  # 标记是否遇到过占位段落
    
    #print(f"\n=== 开始提取脚注 ===")
    #print(f"标题: {paragraphs[start_index].text.strip()}")
    #print(f"起始段落索引: {start_para_idx}")
    #print(f"结束段落索引: {end_index}")
    
    while current_idx < end_index and len(footnote_lines) < max_lines:
        # 定期检查取消状态
        if cancellation_token and cancellation_token.is_cancelled():
            print("脚注提取被用户取消")
            return "", "", "已取消"
        
        paragraph = paragraphs[current_idx]
        text = paragraph.text.strip()
        
        #print(f"检查段落 {current_idx}: '{text}'")
        
        # 检查终止条件
        # 章节名匹配终止条件（精确匹配目录中的章节名）
        if toc_cleaned_titles and text in toc_cleaned_titles:
            #print(f"  -> 终止原因: 匹配目录章节名: {text}")
            break
            
        #if TABLE_NUMBER_PATTERN.match(text):  # 表格编号格式的新标题
            #print(f"  -> 终止原因: 匹配TABLE_NUMBER_PATTERN: {text}")
            #break
            
        if is_fully_italic(paragraph):  # 全斜体段落
            #print(f"  -> 终止原因: 全斜体段落")
            break
            
        if contains_custom_keyword(paragraph, custom_keywords):  # 包含关键词（默认+自定义）
            content_before = extract_content_before_custom_keyword(paragraph, custom_keywords)
            if content_before:
                #print(f"  -> 包含关键词，提取前缀: '{content_before}'")
                footnote_lines.append(content_before)
                processed_text = process_superscript_subscript_text(paragraph)
                # 对处理后的文本再次应用关键词提取逻辑
                content_before_processed = extract_content_before_custom_keyword_from_text(processed_text, custom_keywords)
                processed_footnote_lines.append(content_before_processed)
            #print(f"  -> 终止原因: 包含关键词")
            break
        
        # 检查是否为占位段落
        if is_placeholder_paragraph(paragraph):
            #print(f"  -> 检测到占位段落: '{text}'")
            # 如果这是第一次遇到占位段落，清空已收集的内容
            if not found_placeholder:
                #print(f"  -> 清空之前收集的脚注内容 ({len(footnote_lines)} 行)")
                footnote_lines.clear()
                processed_footnote_lines.clear()
                found_placeholder = True
            
            # 跳过占位段落，从下一个段落开始重新收集脚注
            #print(f"  -> 跳过占位段落，从下一段开始收集")
            current_idx += 1
            continue
            
        # 添加到脚注
        if text:
            #print(f"  -> 添加到脚注: '{text}'")
            footnote_lines.append(text)
            processed_text = process_superscript_subscript_text(paragraph)
            processed_footnote_lines.append(processed_text)
        #else:
            #print(f"  -> 跳过空段落")
        
        current_idx += 1
    
    # 构造返回结果
    footnote_text = "\n".join(footnote_lines)
    processed_footnote_text = "\n".join(processed_footnote_lines)
    status = "成功" if footnote_lines else "脚注未找到"
    
    #print(f"=== 脚注提取完成 ===")
    #print(f"原始脚注文本:\n{footnote_text}")
    #print(f"处理后脚注文本:\n{processed_footnote_text}")
    print(f"状态: {status}")
    #if found_placeholder:
        #print(f"注意: 处理过程中遇到了占位段落，已重置脚注收集")
    print(f"")
    
    return footnote_text, processed_footnote_text, status