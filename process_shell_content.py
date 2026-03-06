#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于预提取的数据生成 Shell 文档（只包含标题和脚注）
删除其他所有内容（包括表格、图片等）
保留原始段落格式（字体、上标、下标等）。
采用与参考程序完全一致的脚注定位逻辑（基于表格、图片、占位段落四级策略）

注意：此脚本必须由 main.py 调用，使用已提取的标题和脚注数据
"""

import os
import sys
import shutil
import re
from docx import Document
from typing import List, Dict, Any, Optional
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 将脚本所在目录加入模块搜索路径，确保能导入同一目录下的参考模块
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from enhanced_footnote_extractor import (
    is_placeholder_paragraph,
    collect_elements_in_range,
    find_footnote_start_position
)
from utils import is_fully_italic, contains_custom_keyword, extract_content_before_custom_keyword


def set_document_default_font(doc, font_name='Times New Roman', chinese_font='宋体', font_size=12):
    """
    设置文档的默认字体和段落格式
    
    Args:
        doc: Document 对象
        font_name: 西文字体名称，默认 Times New Roman
        chinese_font: 中文字体名称，默认宋体
        font_size: 字体大小（磅），默认 12pt（小四）
    
    Returns:
        doc: 设置后的 Document 对象
    """
    # 获取默认样式
    style = doc.styles['Normal']
    
    # 设置默认字体
    font = style.font
    font.name = font_name
    font.size = Pt(font_size)
    
    # 设置中文字体（如果需要）
    style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), chinese_font)
    
    # 设置段落格式
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)  # 段前 0 磅
    paragraph_format.space_after = Pt(0)   # 段后 0 磅
    paragraph_format.line_spacing = 1.0     # 单倍行距
    
    return doc


def format_text_with_special_codes(text: str):
    """
    处理文本中的特殊标记代码，返回格式化后的文本和格式信息
    
    支持的格式：
    - (*ESC*){super XX} -> 只保留 XX，设置为上标
    - (*ESC*){sub XX} -> 只保留 XX，设置为下标
    - (*ESC*){unicode XXXX} -> 替换为对应的 Unicode 字符
    
    Returns:
        list: 包含 (文本片段，是否上标，是否下标) 的元组列表
    """
    if not text:
        return []
    
    result = []
    
    # 定义正则表达式模式
    super_pattern = r'\(\*ESC\*\)\{super\s+([^}]+)\}'
    sub_pattern = r'\(\*ESC\*\)\{sub\s+([^}]+)\}'
    unicode_pattern = r'\(\*ESC\*\)\{unicode\s+([0-9A-Fa-f]+)\}'
    
    # 先处理 unicode 替换
    def replace_unicode(match):
        hex_value = match.group(1)
        try:
            char = chr(int(hex_value, 16))
            return char
        except (ValueError, OverflowError):
            return match.group(0)  # 如果转换失败，保持原样
    
    text = re.sub(unicode_pattern, replace_unicode, text)
    
    # 分割文本并处理上标/下标
    # 我们需要找到所有匹配的位置，然后按位置分割文本
    matches = []
    
    for match in re.finditer(super_pattern, text):
        matches.append({
            'start': match.start(),
            'end': match.end(),
            'type': 'super',
            'content': match.group(1)
        })
    
    for match in re.finditer(sub_pattern, text):
        matches.append({
            'start': match.start(),
            'end': match.end(),
            'type': 'sub',
            'content': match.group(1)
        })
    
    # 按位置排序
    matches.sort(key=lambda x: x['start'])
    
    # 如果没有找到任何特殊标记，返回普通文本
    if not matches:
        return [(text, False, False)]
    
    # 构建结果列表
    current_pos = 0
    for match_info in matches:
        # 添加标记前的普通文本
        if match_info['start'] > current_pos:
            normal_text = text[current_pos:match_info['start']]
            result.append((normal_text, False, False))
        
        # 添加特殊格式的文本
        content = match_info['content']
        if match_info['type'] == 'super':
            result.append((content, True, False))
        else:  # sub
            result.append((content, False, True))
        
        current_pos = match_info['end']
    
    # 添加最后的普通文本
    if current_pos < len(text):
        normal_text = text[current_pos:]
        result.append((normal_text, False, False))
    
    return result


def add_formatted_paragraph(doc, text: str, style='Normal'):
    """
    添加一个格式化的段落到文档，处理特殊标记代码
    对于标题设置大纲级别（Heading 1）并应用默认字体（Times New Roman/宋体，12pt）
    对于正文（Normal），应用默认字体设置
    
    Args:
        doc: Document 对象
        text: 要添加的文本（可能包含特殊标记）
        style: 段落样式（'Heading 1' 或 'Normal'）
    """
    para = doc.add_paragraph(style=style)
    
    if not text:
        return para
    
    # 处理文本中的特殊代码
    formatted_parts = format_text_with_special_codes(text)
    
    # 如果是 Heading 1 样式，需要显式设置字体以覆盖默认的 Heading 样式
    is_heading = (style == 'Heading 1')
    
    # 添加每个部分
    for part_text, is_super, is_sub in formatted_parts:
        run = para.add_run(part_text)
        
        # 如果是标题，显式设置字体为默认值（覆盖 Heading 1 的默认样式）
        if is_heading:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # 设置中文字体
            run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
            # 确保文本不加粗
            run.font.bold = False
        
        # 设置上标/下标
        if is_super:
            run.font.superscript = True
        elif is_sub:
            run.font.subscript = True
    
    return para


def process_shell_file(
    auxiliary_data: List[Dict[str, Any]],
    source_path: str,
    output_path: Optional[str] = None,
    project_id: Optional[str] = None
):
    """
    基于 auxiliary_data 生成 shell 文档（只包含标题和脚注）
    
    Args:
        auxiliary_data: 来自主函数的辅助数据列表，每个元素包含：
                       - 文件名称：Output Name
                       - 表格名称：标题
                       - processed_title: 处理后的标题（带格式标记）
                       - processed_footnote: 处理后的脚注（带格式标记）
        source_path: 源文件路径（用于确定输出目录）
        output_path: 输出文件路径（可选，默认在源文件同目录下）
        project_id: 项目编号（可选，用于输出文件名前缀）
    
    输出格式：
    - 标题：大纲级别为 1 级（Heading 1）
    - 脚注：大纲级别为正文（Normal）
    - 支持特殊格式标记：
      - (*ESC*){super XX} -> XX 显示为上标
      - (*ESC*){sub XX} -> XX 显示为下标
      - (*ESC*){unicode XXXX} -> 替换为对应 Unicode 字符
    """
    # 确定输出路径
    if output_path is None:
        target_dir = os.path.dirname(source_path)
        if project_id:
            output_path = os.path.join(target_dir, f"{project_id}_title_footnote_shell.docx")
        else:
            output_path = os.path.join(target_dir, "title_footnote_shell.docx")
    
    print(f"[1/2] 准备生成 shell 文档：{output_path}")
    
    # 验证必要参数
    if not auxiliary_data:
        raise ValueError("错误：auxiliary_data 不能为空")
    
    print("[2/2] 使用 auxiliary_data 生成 shell 文档...")
    
    # 创建新文档并设置默认字体
    new_doc = Document()
    set_document_default_font(new_doc, font_name='Times New Roman', chinese_font='宋体', font_size=12)
    
    # 遍历 auxiliary_data，添加标题和脚注
    for idx, item in enumerate(auxiliary_data, 1):
        # 获取处理后的标题和脚注
        # 注意：需要从原始 results 数据中获取 processed_title 和 processed_footnote
        # auxiliary_data 只包含基本信息，需要通过某种方式传递完整数据
        
        # 从 item 中提取信息
        file_name = item.get('文件名称', '')
        table_name = item.get('表格名称', '')
        
        # 添加标题（大纲级别为 1 级，但不使用 Heading 1 样式以保留默认字体设置）
        # 优先使用 processed_title，如果没有则使用 table_name
        title_text = item.get('processed_title', table_name)
        if title_text:
            add_formatted_paragraph(new_doc, title_text, style='Heading 1')
        
        # 添加脚注（Normal）
        footnote_text = item.get('processed_footnote', '')
        if footnote_text:
            add_formatted_paragraph(new_doc, footnote_text, style='Normal')
    
    # 保存文档
    new_doc.save(output_path)
    print(f"✅ Shell 文档已生成：{output_path}")
    print(f"   共处理 {len(auxiliary_data)} 个条目")


if __name__ == "__main__":
    print("错误：此脚本不能独立运行，必须由 main.py 调用")
    print("用法：在 main.py 中导入并调用 process_shell_file() 函数")
    sys.exit(1)